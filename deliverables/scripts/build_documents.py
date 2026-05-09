from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
    Image as RLImage,
    KeepTogether,
)
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SHOTS = ROOT / "screenshots" / "annotated"
OUT.mkdir(parents=True, exist_ok=True)

TODAY = "May 2026"
BLUE = "2563EB"
NAVY = "0F172A"
TEAL = "14B8A6"
ORANGE = "F97316"
LIGHT_BLUE = "EFF6FF"
LIGHT_LINE = "D7E3F8"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = NAVY) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.3)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyper_clean_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].font.color.rgb = RGBColor.from_string("334155")

    for style_name, size, color in [
        ("Title", 26, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, NAVY),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def set_margins(section) -> None:
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)


def add_doc_cover(doc: Document, title: str, subtitle: str, kicker: str, doc_type: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("475569")

    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Document", doc_type),
        ("Prepared for", "Pharmacy Owners, Directors, Managers, and Operations Teams"),
        ("System", "Smart Pharmacy Management System"),
        ("Date", TODAY),
    ]
    for i, (left, right) in enumerate(rows):
        set_cell_shading(table.rows[i].cells[0], LIGHT_BLUE)
        set_cell_text(table.rows[i].cells[0], left, True, BLUE)
        set_cell_text(table.rows[i].cells[1], right)
    doc.add_page_break()


def add_doc_section(doc: Document, title: str, paragraphs=None, bullets=None, table=None) -> None:
    doc.add_heading(title, level=1)
    for para in paragraphs or []:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.12
    for bullet in bullets or []:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run("- ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(bullet)
    if table:
        headers, rows = table
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            set_cell_shading(t.rows[0].cells[j], BLUE)
            set_cell_text(t.rows[0].cells[j], h, True, "FFFFFF")
        for row in rows:
            cells = t.add_row().cells
            for j, val in enumerate(row):
                set_cell_text(cells[j], val)
        doc.add_paragraph()


proposal_sections = [
    {
        "title": "Cover Letter",
        "paragraphs": [
            "To: Pharmacy Owners, Managing Directors, Branch Managers, and Decision Makers",
            "Subject: Proposal for the Implementation of the Smart Pharmacy Management System",
            "We respectfully submit this proposal for the implementation of Smart Pharmacy, a secure web-based operations platform designed to bring point-of-sale, stock control, purchasing, sales returns, expenses, daily closing, reporting, alerts, approvals, and audit visibility into one coordinated pharmacy system.",
            "The solution responds to a practical reality: many pharmacies sell quickly, but they lose visibility when stock, expiry, cash, purchasing, and branch accountability are handled in separate books, spreadsheets, or informal messages. Smart Pharmacy closes that gap by making every important business movement recorded, permission-controlled, reportable, and traceable.",
            "We recommend Smart Pharmacy as a strong operational foundation for pharmacies that want faster service, safer stock handling, better owner visibility, and scalable digital growth.",
        ],
    },
    {
        "title": "1. Executive Summary",
        "paragraphs": [
            "Smart Pharmacy is a comprehensive pharmacy operations platform built around fast POS selling, accurate inventory control, branch-level accountability, controlled approvals, real-time alerts, and management reporting.",
            "The system is not only a cashier screen. It is an operating control layer for the pharmacy: sales affect stock, purchases create inventory batches, returns affect revenue and stock only when approved, expenses affect daily closing, alerts identify risk, and reports give management evidence for decisions.",
        ],
        "bullets": [
            "Fast single-screen POS for product search, cart, payment, receipt, and POS expenses.",
            "Inventory tracked by branch, product, batch, expiry date, available quantity, and cost.",
            "Purchases, transfers, adjustments, returns, and daily closing controlled through statuses and approvals.",
            "Inventory alerts for low stock, out-of-stock, expiring soon, and expired stock.",
            "Role-based access for Owner, Admin, Pharmacist, Cashier, and Storekeeper.",
            "Responsive online interface usable across desktop, tablet, and mobile screens.",
            "Future-ready foundation for AI assistance and SMS, email, and WhatsApp communication.",
        ],
    },
    {
        "title": "2. Business Problem and Discovery",
        "paragraphs": [
            "The discovery behind the system is simple: pharmacy operations are connected, but manual tools often treat them as separate tasks. A cashier sells, a storekeeper receives stock, an owner checks cash, a manager watches expiry, and a pharmacist handles returns or prescriptions. If these actions are not connected, the business becomes difficult to control.",
            "Common operational challenges include stock not matching sales, expired products remaining sellable, unclear cash differences, delayed approval of returns and transfers, weak branch accountability, and limited visibility for owners who are not physically present in the pharmacy.",
        ],
        "table": (
            ["Current Challenge", "Business Risk", "Smart Pharmacy Response"],
            [
                ["Manual stock records", "Stock mismatch, losses, and reorder delays", "Inventory updates from purchases, sales, returns, adjustments, and transfers"],
                ["Uncontrolled expiry handling", "Expired medicine can remain in circulation", "Expiry alerts and automatic write-off workflow"],
                ["Disconnected cashier closing", "Cash differences are hard to prove", "Daily closing compares expected cash against counted cash"],
                ["Informal approvals", "Sensitive actions happen without accountability", "Role permissions, approval statuses, and notifications"],
                ["Limited owner visibility", "Management decisions depend on delayed reports", "Dashboard, reports, alerts, and audit logs"],
            ],
        ),
    },
    {
        "title": "3. Proposed Solution Overview",
        "paragraphs": [
            "Smart Pharmacy operates as a secure, role-based web platform. Each user sees only the modules and actions allowed by their role. The same system shell connects sales, stock, finance, and reporting so that one action automatically informs the next.",
            "The design principle is to make daily work fast for frontline users while keeping the business controlled for owners and administrators.",
        ],
        "bullets": [
            "POS completes sales, updates stock, supports receipts, and records POS expenses.",
            "Product setup manages product types, categories, units, package structure, and prices.",
            "Purchasing brings stock into inventory batches with supplier and cost records.",
            "Inventory monitors available quantity, expiry status, batch movement, and stock value.",
            "Sales returns, adjustments, transfers, and daily closing follow controlled approval workflows.",
            "Reports turn operational records into sales, stock, purchase, expense, and profit insight.",
            "Topbar counters and notifications keep responsible users aware of urgent actions.",
        ],
    },
    {
        "title": "4. Strategic Objectives",
        "bullets": [
            "Improve sales speed and reduce cashier friction through a single-screen POS.",
            "Protect stock accuracy using unit conversion, base-unit logic, batches, and movement records.",
            "Reduce expired-stock risk through alerts and automatic write-off readiness.",
            "Strengthen cash accountability through daily closing and expense integration.",
            "Create branch-level visibility for owners and managers.",
            "Standardize approvals for sensitive business movements.",
            "Prepare the pharmacy for digital communication, AI assistance, and future integrations.",
        ],
    },
    {
        "title": "5. Functional Modules",
        "table": (
            ["Module", "Purpose", "Primary Users"],
            [
                ["Dashboard", "Shows sales, profit, expenses, trends, top products, and activity", "Owner, Admin, Managers"],
                ["Smart POS", "Fast sales, receipt printing, payment selection, today sales, POS expenses", "Cashier, Pharmacist, Owner"],
                ["Product Setup", "Product types, categories, units, package structure, and price rules", "Admin, Storekeeper"],
                ["Suppliers and Purchases", "Supplier records, purchase items, receiving, and purchase reports", "Storekeeper, Admin"],
                ["Inventory Control", "Batch stock, expiry, movements, low stock, and current availability", "Storekeeper, Pharmacist, Owner"],
                ["Sales Returns", "Customer returns with condition, refund, approval, and stock rules", "Cashier, Pharmacist, Owner"],
                ["Expenses and Daily Closing", "Expense capture, expected cash, counted cash, and verification", "Cashier, Admin, Owner"],
                ["Reports", "Sales, stock, purchases, profit, expenses, and prescriptions", "Owner, Admin, Managers"],
                ["Users, Roles, Audit", "Access control, permissions, activity logs, and accountability", "Owner, Admin"],
            ],
        ),
    },
    {
        "title": "6. Workflow Control Model",
        "paragraphs": [
            "The system connects stock and money movement through one controlled workflow:",
            "Setup -> Products and Units -> Suppliers -> Purchases -> Inventory Batches -> POS Sales -> Sales Returns -> Expenses -> Daily Closing -> Reports -> Dashboard -> Alerts and Notifications.",
            "This creates a full operational chain where each business event has a record, a responsible user, a branch, and a reportable effect.",
        ],
    },
    {
        "title": "7. Why Smart Pharmacy Is Advised Over Ordinary Alternatives",
        "table": (
            ["Decision Factor", "Ordinary POS or Spreadsheet", "Smart Pharmacy Advantage"],
            [
                ["Stock accuracy", "Often separate from sales or manually reconciled", "Sales, purchases, returns, transfers, and adjustments update stock records"],
                ["Expiry safety", "Depends on manual checking", "Expiry alerts and automatic write-off workflow protect sellable stock"],
                ["Cash control", "Cashier totals may be manually compared", "Daily closing calculates expected cash from sales, refunds, and expenses"],
                ["Branch control", "Branch visibility is usually limited", "Branch-aware inventory, sales, purchases, reports, and users"],
                ["Approvals", "Informal approvals through calls or messages", "Permission-based approval buttons and notification targeting"],
                ["Online access", "Often tied to one computer", "Responsive web access across desktop, tablet, and mobile screens"],
                ["Growth readiness", "Hard to extend", "Built for AI, internal communication, messaging, and integrations"],
            ],
        ),
    },
    {
        "title": "8. Online Responsiveness and Multi-Screen Access",
        "paragraphs": [
            "The interface is responsive across common working screens. A manager can review dashboard performance on a desktop, a cashier can use POS on a large counter display, and authorized users can still review core information from smaller screens.",
            "This matters because pharmacy decisions do not always happen at the cashier desk. Owners and managers often need to view sales, stock risks, approvals, and reports while away from the branch.",
        ],
    },
    {
        "title": "9. Security, Roles, and Accountability",
        "paragraphs": [
            "Smart Pharmacy uses role-based access control so that the system can separate normal operation from sensitive approval power. Cashiers can sell and submit closings, while owners or administrators can approve returns, transfers, stock adjustments, and daily closings.",
            "Activity logs and notification records support accountability. This makes it easier to answer who performed an action, when it happened, and which record was affected.",
        ],
    },
    {
        "title": "10. Future Improvement Roadmap",
        "bullets": [
            "AI assistant for product lookup, stock risk explanation, sales insights, and management summaries.",
            "Internal communication hub for task notes, approvals, and branch-to-branch operational messages.",
            "SMS, email, and WhatsApp notifications for stock alerts, approvals, customer receipts, supplier updates, and owner summaries.",
            "Predictive inventory recommendations based on sales velocity, expiry risk, and reorder patterns.",
            "Prescription workflow expansion, including upload, pharmacist review, and sale linkage.",
            "Supplier portals and customer loyalty or refill reminders.",
            "Mobile app layer for owners, branch supervisors, and field approvals.",
        ],
    },
    {
        "title": "11. Implementation Approach",
        "table": (
            ["Phase", "Focus", "Outcome"],
            [
                ["Phase 1", "Setup, users, roles, products, units, prices, branches", "System foundation ready for controlled use"],
                ["Phase 2", "POS, purchases, inventory, receipts, expenses", "Daily operations move into the platform"],
                ["Phase 3", "Returns, adjustments, transfers, daily closing, alerts", "Control and approval workflows become active"],
                ["Phase 4", "Reports, dashboard, audit review, optimization", "Management insight and accountability improve"],
                ["Phase 5", "AI, SMS, email, WhatsApp, integrations", "Advanced communication and intelligence layer"],
            ],
        ),
    },
    {
        "title": "12. Expected Benefits",
        "bullets": [
            "Faster customer service and cleaner cashier workflow.",
            "Better stock accuracy and fewer inventory surprises.",
            "Reduced expired-stock risk and stronger medicine safety.",
            "Clearer cash accountability at daily closing.",
            "More reliable branch and user accountability.",
            "Better management decisions using reports and dashboards.",
            "A scalable platform that can grow beyond the MVP without replacing the core system.",
        ],
    },
    {
        "title": "13. Recommendation",
        "paragraphs": [
            "Smart Pharmacy is recommended as a practical and future-ready operating system for pharmacies that want more than a basic POS. It brings together sales speed, stock control, cash control, alerts, approvals, reporting, and auditability.",
            "The strongest value is that the system turns pharmacy work into one connected digital workflow: accurate stock, fast sales, controlled approvals, expiry safety, owner visibility, and branch accountability.",
        ],
        "table": (
            ["Decision Area", "Recommended Next Step"],
            [
                ["Executive approval", "Approve Smart Pharmacy as the preferred digital operating platform for the pharmacy."],
                ["Implementation scope", "Begin with setup, users, products, POS, purchases, inventory, expenses, daily closing, and reports."],
                ["Operational readiness", "Nominate Owner/Admin, Cashier, Pharmacist, and Storekeeper users for initial training."],
                ["Future extension", "Plan AI assistance and SMS, email, and WhatsApp communication as the next enhancement layer."],
            ],
        ),
    },
]


manual_modules = [
    ("01-login", "Login and Secure Access", "Use the login page to access the system with email or username and password.", ["Enter email or username.", "Enter password.", "Click Log in.", "The visible menus are loaded according to the user's role."]),
    ("02-dashboard", "Dashboard", "Use the dashboard to monitor sales, profit, expenses, trends, top products, and activity.", ["Select branch and date range.", "Click Apply.", "Review revenue, profit, trend, and product performance.", "Open reports or alerts for deeper investigation."]),
    ("03-smart-pos", "Smart POS", "Use POS to search medicines, add items to cart, select payment method, complete sale, and print receipts.", ["Search product by name, category, type, code, or barcode.", "Choose retail or wholesale mode.", "Add product to current sale.", "Enter discount or customer details if needed.", "Select payment method and complete sale."]),
    ("04-pharmacy-setup", "Pharmacy Setup", "Maintain pharmacy profile, branches, selling rules, expiry settings, and receipt settings.", ["Open Pharmacy Setup.", "Update profile and branch information.", "Configure currency, expiry warning days, and receipt footer.", "Save changes."]),
    ("05-product-setup", "Product Setup", "Manage product types, categories, units, package structure, product records, and price rules.", ["Create or edit product types.", "Create categories and units.", "Add products with generic name, strength, brand, and base unit.", "Configure package structure and retail/wholesale prices.", "Use import or export where needed."]),
    ("06-suppliers", "Suppliers", "Store supplier contacts used during purchases.", ["Add supplier name and contact details.", "Activate or deactivate suppliers.", "Select supplier when creating a purchase."]),
    ("07-purchases", "Purchases", "Record purchases and receive stock into inventory batches.", ["Create purchase.", "Select supplier and branch.", "Add purchase items with unit, batch, expiry, quantity, and cost.", "Receive purchase to increase inventory."]),
    ("08-inventory", "Current Inventory", "Track available stock by branch, product, batch, expiry, quantity, cost, and status.", ["Use filters for branch, status, expiry, and low stock.", "Review available quantity and expiry date.", "Open movements for traceability.", "Block, adjust, or mark stock according to permission."]),
    ("09-inventory-movements", "Inventory Movements", "Review the full history of stock increases, reductions, adjustments, transfers, sales, and returns.", ["Filter by product, branch, movement type, or date.", "Review before and after balances.", "Use movement history for investigation and audit."]),
    ("10-inventory-alerts", "Inventory Alerts", "Monitor low stock, out-of-stock, expiring soon, and expired inventory alerts.", ["Open Inventory Alerts.", "Filter by alert type or status.", "Generate alerts manually if needed.", "Mark read, resolve, or ignore according to permission."]),
    ("11-stock-adjustments", "Stock Adjustments", "Correct stock for damage, expiry, loss, physical count difference, found stock, or general correction.", ["Create adjustment request.", "Select branch and inventory batch.", "Choose direction and quantity.", "Submit for approval.", "Owner/Admin approves or rejects."]),
    ("12-stock-transfers", "Stock Transfers", "Move stock between branches through approval, dispatch, and receiving stages.", ["Create transfer from source to destination branch.", "Search inventory batch and select unit/quantity.", "Save draft for approval.", "Dispatch after approval.", "Receive at destination branch."]),
    ("13-sales-history", "Sales History", "View completed sales, receipt information, payment methods, totals, and profit effect.", ["Filter sales by branch, date, payment, or status.", "Open receipt details.", "Reprint receipts or cancel sale where permission allows."]),
    ("14-sales-returns", "Sales Returns", "Manage returned items with condition, refund method, approval, and stock restoration rules.", ["Search original sale.", "Select items and return quantity.", "Set condition as sellable, opened, damaged, or expired.", "Submit return request.", "Owner/Admin approves or rejects."]),
    ("15-expenses", "Expenses", "Record business expenses and categories for daily cash accountability and reporting.", ["Create or manage expense categories.", "Record title, category, amount, date, and payment method.", "Void or edit according to permission.", "Review expense reports."]),
    ("16-daily-closing", "Daily Closing", "Compare expected cash against counted cash for cashier and branch accountability.", ["Select branch and date.", "Calculate expected sales, refunds, and expenses.", "Enter counted cash.", "Save draft or submit.", "Owner/Admin verifies, rejects, or recalculates."]),
    ("17-report-center", "Report Center", "Open sales, stock, purchase, profit, expense, and prescription reports.", ["Choose the required report.", "Apply branch and date filters.", "Review summary and detailed tables.", "Export where needed."]),
    ("18-profit-report", "Profit Report", "Analyze net sales, cost sold, gross profit, expenses, returns, and net profit.", ["Select date range and branch.", "Review sales and cost figures.", "Compare gross and net profit.", "Use the report for owner decisions."]),
    ("19-users", "Users", "Create and manage user accounts, branch assignment, role assignment, status, and password resets.", ["Add a staff user.", "Assign branch and role.", "Activate or deactivate account.", "Reset password when required."]),
    ("20-roles", "Roles and Permissions", "Control what each user type can view, create, approve, reject, verify, and report.", ["Open Roles.", "Review permission groups.", "Update permissions for the role.", "Keep sensitive approvals limited to trusted users."]),
    ("21-activity-logs", "Activity Logs", "Audit important actions performed in the system.", ["Filter by log name, user, date, or action.", "Review record details.", "Use logs to investigate changes and approvals."]),
    ("22-quick-search", "Quick Search", "Find products, receipts, batches, purchases, returns, transfers, adjustments, and expenses from the topbar.", ["Click the topbar search field.", "Type at least two characters.", "Review permission-aware results.", "Open the selected record."]),
    ("23-mobile-dashboard", "Mobile Dashboard", "Review dashboard metrics on a phone-sized screen.", ["Use filters vertically.", "Scroll through metric cards and trends.", "Open alerts, messages, or profile controls from the compact topbar."]),
    ("24-mobile-pos", "Mobile POS", "Use the POS workflow from smaller screens where operationally appropriate.", ["Search medicine.", "Add items.", "Review cart and payment controls.", "Complete sale if the role and device workflow allow it."]),
]


def build_proposal_docx(path: Path) -> None:
    doc = Document()
    set_margins(doc.sections[0])
    add_hyper_clean_styles(doc)
    add_doc_cover(
        doc,
        "Proposal for the Smart Pharmacy Management System",
        "A Comprehensive Digital Pharmacy Operations, POS, Inventory, Finance, Reporting, and Control Platform",
        "Formal Executive Proposal",
        "Executive Proposal",
    )
    for section in proposal_sections:
        add_doc_section(doc, **section)
    doc.save(path)


def build_manual_docx(path: Path) -> None:
    doc = Document()
    set_margins(doc.sections[0])
    add_hyper_clean_styles(doc)
    add_doc_cover(
        doc,
        "Smart Pharmacy User Manual",
        "Operational Guide with Screen-by-Screen Module Walkthroughs and Callout Annotations",
        "System Manual",
        "User Manual",
    )
    add_doc_section(
        doc,
        "How to Use This Manual",
        paragraphs=[
            "This manual explains the main Smart Pharmacy modules using the actual system screens. Each screenshot includes callout arrows that point to the controls or areas users should notice first.",
            "Role permissions may change which menus and buttons are visible. If a user cannot see a module or action, the administrator should review the assigned role and permission set.",
        ],
        bullets=[
            "Owner/Admin: broad system setup, approvals, reports, users, and audit review.",
            "Cashier: POS, receipts, POS expenses, today sales, returns request, and daily closing submission.",
            "Pharmacist: POS, prescriptions, stock visibility, return requests, adjustments, transfers, and closing submission.",
            "Storekeeper: products, suppliers, purchases, inventory, adjustments, transfers, and stock reports.",
        ],
    )
    for slug, title, purpose, steps in manual_modules:
        doc.add_page_break()
        doc.add_heading(title, level=1)
        doc.add_paragraph(purpose)
        doc.add_heading("Main Steps", level=2)
        for step in steps:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            r = p.add_run("- ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(BLUE)
            p.add_run(step)
        img = SHOTS / f"{slug}.png"
        if img.exists():
            doc.add_heading("Annotated Screen", level=2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width = Inches(2.25) if "mobile" in slug else Inches(6.8)
            p.add_run().add_picture(str(img), width=width)
            cap = doc.add_paragraph(f"Figure: {title} module screen with callout arrows.")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(8.5)
            cap.runs[0].font.color.rgb = RGBColor.from_string("64748B")
    doc.save(path)


styles = getSampleStyleSheet()
pdf_title = ParagraphStyle(
    "PdfTitle",
    parent=styles["Title"],
    fontName="Helvetica-Bold",
    fontSize=24,
    leading=29,
    textColor=colors.HexColor(f"#{NAVY}"),
    alignment=TA_CENTER,
    spaceAfter=14,
)
pdf_subtitle = ParagraphStyle(
    "PdfSubtitle",
    parent=styles["BodyText"],
    fontSize=12,
    leading=16,
    textColor=colors.HexColor("#475569"),
    alignment=TA_CENTER,
    spaceAfter=18,
)
pdf_h1 = ParagraphStyle(
    "PdfH1",
    parent=styles["Heading1"],
    fontName="Helvetica-Bold",
    fontSize=15,
    leading=20,
    textColor=colors.HexColor(f"#{NAVY}"),
    spaceBefore=8,
    spaceAfter=8,
)
pdf_body = ParagraphStyle(
    "PdfBody",
    parent=styles["BodyText"],
    fontSize=9.6,
    leading=13,
    textColor=colors.HexColor("#334155"),
    spaceAfter=6,
)
pdf_bullet = ParagraphStyle(
    "PdfBullet",
    parent=pdf_body,
    leftIndent=13,
    firstLineIndent=-8,
)
pdf_caption = ParagraphStyle(
    "PdfCaption",
    parent=styles["BodyText"],
    fontSize=8,
    leading=10,
    textColor=colors.HexColor("#64748B"),
    alignment=TA_CENTER,
)


def pdf_cover(story, title: str, subtitle: str, doc_type: str) -> None:
    story.append(Spacer(1, 0.85 * inch))
    story.append(Paragraph(title, pdf_title))
    story.append(Paragraph(subtitle, pdf_subtitle))
    data = [
        ["Document", doc_type],
        ["Prepared for", "Pharmacy Owners, Directors, Managers, and Operations Teams"],
        ["System", "Smart Pharmacy Management System"],
        ["Date", TODAY],
    ]
    t = Table(data, colWidths=[1.55 * inch, 4.85 * inch])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{LIGHT_BLUE}")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor(f"#{BLUE}")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LIGHT_LINE}")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(t)
    story.append(PageBreak())


def add_pdf_section(story, title: str, paragraphs=None, bullets=None, table=None) -> None:
    story.append(Paragraph(title, pdf_h1))
    for para in paragraphs or []:
        story.append(Paragraph(para, pdf_body))
    for bullet in bullets or []:
        story.append(Paragraph(f"- {bullet}", pdf_bullet))
    if table:
        headers, rows = table
        data = [headers] + rows
        col_width = 6.8 * inch / len(headers)
        t = Table(data, colWidths=[col_width] * len(headers), repeatRows=1)
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BLUE}")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                    ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{LIGHT_LINE}")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.append(Spacer(1, 4))
        story.append(t)
        story.append(Spacer(1, 8))


def rl_image(path: Path, max_w: float, max_h: float) -> RLImage:
    with Image.open(path) as im:
        w, h = im.size
    scale = min(max_w / w, max_h / h)
    return RLImage(str(path), width=w * scale, height=h * scale)


def build_proposal_pdf(path: Path) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.5 * inch, bottomMargin=0.48 * inch)
    story = []
    pdf_cover(
        story,
        "Proposal for the Smart Pharmacy Management System",
        "A Comprehensive Digital Pharmacy Operations, POS, Inventory, Finance, Reporting, and Control Platform",
        "Executive Proposal",
    )
    for section in proposal_sections:
        add_pdf_section(story, **section)
    doc.build(story)


def build_manual_pdf(path: Path) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=0.42 * inch, leftMargin=0.42 * inch, topMargin=0.45 * inch, bottomMargin=0.45 * inch)
    story = []
    pdf_cover(
        story,
        "Smart Pharmacy User Manual",
        "Operational Guide with Screen-by-Screen Module Walkthroughs and Callout Annotations",
        "User Manual",
    )
    add_pdf_section(
        story,
        "How to Use This Manual",
        paragraphs=[
            "This manual explains the main Smart Pharmacy modules using actual system screens. Each screenshot includes callout arrows that point to the controls or areas users should notice first.",
            "Role permissions may change visible menus and buttons. If a user cannot see a module or action, the administrator should review the assigned role and permission set.",
        ],
        bullets=[
            "Owner/Admin: broad setup, approvals, reports, users, and audit review.",
            "Cashier: POS, receipts, expenses, today sales, returns request, and daily closing submission.",
            "Pharmacist: POS, prescriptions, stock visibility, return requests, adjustments, transfers, and closing submission.",
            "Storekeeper: products, suppliers, purchases, inventory, adjustments, transfers, and stock reports.",
        ],
    )
    for slug, title, purpose, steps in manual_modules:
        story.append(PageBreak())
        items = [Paragraph(title, pdf_h1), Paragraph(purpose, pdf_body)]
        items.append(Paragraph("Main Steps", ParagraphStyle("MiniHead", parent=pdf_body, fontName="Helvetica-Bold", textColor=colors.HexColor(f"#{BLUE}"))))
        for step in steps:
            items.append(Paragraph(f"- {step}", pdf_bullet))
        img = SHOTS / f"{slug}.png"
        if img.exists():
            items.append(Spacer(1, 6))
            max_w = 2.9 * inch if "mobile" in slug else 7.25 * inch
            max_h = 5.3 * inch if "mobile" in slug else 4.6 * inch
            items.append(rl_image(img, max_w, max_h))
            items.append(Paragraph(f"Figure: {title} module screen with callout arrows.", pdf_caption))
        story.append(KeepTogether(items))
    doc.build(story)


if __name__ == "__main__":
    build_proposal_docx(OUT / "Smart_Pharmacy_Executive_Proposal.docx")
    build_proposal_pdf(OUT / "Smart_Pharmacy_Executive_Proposal.pdf")
    build_manual_docx(OUT / "Smart_Pharmacy_User_Manual.docx")
    build_manual_pdf(OUT / "Smart_Pharmacy_User_Manual.pdf")
    print("documents built")
